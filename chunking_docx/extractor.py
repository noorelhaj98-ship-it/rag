"""DOCX text extraction."""

from typing import List, Dict, Any


class DOCXExtractor:
    """Extract text from DOCX files using python-docx."""

    def __init__(self):
        self._Document = None

    def _get_document_class(self):
        """Lazy import python-docx."""
        if self._Document is None:
            try:
                from docx import Document
                self._Document = Document
            except ImportError as e:
                raise RuntimeError(
                    "Missing dependency: python-docx. Install with: pip install python-docx"
                ) from e
        return self._Document

    def extract_paragraphs(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX file paragraph by paragraph.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            List of dicts with paragraph_number (1-based) and text
        """
        import re

        Document = self._get_document_class()
        doc = Document(docx_path)
        paragraphs = []

        for i, para in enumerate(doc.paragraphs, start=1):
            text = para.text or ""
            # Normalize line endings and whitespace
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            text = re.sub(r"\n+", "\n", text)
            text = text.strip()

            if text:
                paragraphs.append({"paragraph_number": i, "text": text})

        return paragraphs

    def extract_tables(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract text from tables in DOCX file.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            List of dicts with table content
        """
        Document = self._get_document_class()
        doc = Document(docx_path)
        table_entries = []
        table_count = 0

        # Iterate through all elements in document body
        for element in doc.element.body:
            if element.tag.endswith('tbl'):  # It's a table
                table_count += 1
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break

                if table:
                    table_text_parts = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_text_parts.append(" | ".join(row_text))

                    if table_text_parts:
                        full_table_text = f"\n[Table {table_count}]\n" + "\n".join(table_text_parts)
                        table_entries.append({
                            "paragraph_number": f"table_{table_count}",
                            "text": full_table_text
                        })

        return table_entries
