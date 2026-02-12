import os
import re
import json
import csv
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Dict, Any

# ---- Config ----
SOURCE_FILE = "document.docx"  # Default, can be overridden via command line
OUT_JSONL = "chunks.jsonl"
OUT_CSV = "chunks.csv"

CHUNKING_STRATEGY = "recursive"
TARGET_CHARS = 1200          # target chunk size (characters)
OVERLAP_CHARS = 150          # overlap to preserve context


# ---- DOCX text extraction ----
def extract_paragraphs_text_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract text from DOCX file paragraph by paragraph.
    Returns: list of dicts: [{paragraph_number: 1-based, text: "..."}]
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("Missing dependency: python-docx. Install with: pip install python-docx") from e

    doc = Document(docx_path)
    paragraphs = []
    
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text or ""
        # Normalize line endings and whitespace
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n+", "\n", text)  # collapse multiple newlines
        text = text.strip()
        
        if text:  # Only include non-empty paragraphs
            paragraphs.append({"paragraph_number": i, "text": text})
    
    return paragraphs


def extract_tables_text_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract text from tables in DOCX file.
    Returns: list of dicts with table content.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("Missing dependency: python-docx. Install with: pip install python-docx") from e

    doc = Document(docx_path)
    table_entries = []
    table_count = 0
    
    # Iterate through all elements in document body
    for element in doc.element.body:
        if element.tag.endswith('tbl'):  # It's a table
            table_count += 1
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            
            if table:
                table_text_parts = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text_parts.append(" | ".join(row_text))
                
                if table_text_parts:
                    full_table_text = f"\n[Table {table_count}]\n" + "\n".join(table_text_parts)
                    table_entries.append({
                        "paragraph_number": f"table_{table_count}",
                        "text": full_table_text
                    })
    
    return table_entries


SEPARATORS = [
    "\n\n",   # paragraphs
    "\n",     # lines
    ". ",     # sentences (approx)
    " "       # words
]


def split_recursive(text: str, max_chars: int):
    """
    Recursively split text by increasingly smaller separators until pieces fit.
    Returns a list of pieces (strings) each <= max_chars (best effort).
    """
    text = text.strip()
    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    # Try separators from coarse to fine
    for sep in SEPARATORS:
        if sep in text:
            parts = text.split(sep)
            # Re-add separator when joining later (except last)
            rebuilt = []
            buf = ""
            for idx, p in enumerate(parts):
                if not p:
                    continue
                candidate = (buf + (sep if buf else "") + p).strip() if sep != " " else (buf + (" " if buf else "") + p).strip()
                if len(candidate) <= max_chars:
                    buf = candidate
                else:
                    if buf:
                        rebuilt.append(buf)
                    # if single part too big, recurse
                    if len(p) > max_chars:
                        rebuilt.extend(split_recursive(p, max_chars))
                        buf = ""
                    else:
                        buf = p.strip()
            if buf:
                rebuilt.append(buf)

            # If we actually made progress, return
            if rebuilt and sum(len(x) for x in rebuilt) >= int(0.6 * len(text)):
                return rebuilt

    # Fallback hard split
    return [text[i:i + max_chars].strip() for i in range(0, len(text), max_chars)]


def pack_with_overlap(pieces, overlap_chars: int):
    """
    Pack pieces with overlap between consecutive chunks for context preservation.
    """
    if not pieces:
        return []

    packed = []
    prev = ""
    for p in pieces:
        p = p.strip()
        if not p:
            continue
        if prev and overlap_chars > 0:
            overlap = prev[-overlap_chars:]
            # avoid duplicating if already starts similarly
            if not p.startswith(overlap):
                p = (overlap + "\n" + p).strip()
        packed.append(p)
        prev = p
    return packed


def stable_chunk_id(document_id: str, para_number: int, chunk_index: int, text: str):
    h = hashlib.sha1(f"{document_id}|para{para_number}|c{chunk_index}|{text[:200]}".encode("utf-8")).hexdigest()[:16]
    return f"{document_id}_para{para_number:03d}_c{chunk_index:03d}_{h}"


# ---- Main pipeline ----
def build_chunks(docx_path: str) -> List[Dict[str, Any]]:
    """
    Build chunks from a DOCX file.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    document_id = hashlib.sha1(docx_path.name.encode("utf-8")).hexdigest()[:12]
    created_at = datetime.now(timezone.utc).isoformat()

    # Extract paragraphs and tables
    paragraphs = extract_paragraphs_text_docx(str(docx_path))
    tables = extract_tables_text_docx(str(docx_path))
    
    # Combine all content
    all_content = paragraphs + tables
    all_content.sort(key=lambda x: x["paragraph_number"])

    chunks = []
    global_char_cursor = 0

    for item in all_content:
        para_number = item["paragraph_number"]
        para_text = item["text"] or ""
        para_text_stripped = para_text.strip()

        if not para_text_stripped:
            global_char_cursor += len(para_text)
            continue

        pieces = split_recursive(para_text, TARGET_CHARS)
        pieces = pack_with_overlap(pieces, OVERLAP_CHARS)

        page_search_start = 0
        for idx, chunk_text in enumerate(pieces, start=1):
            normalized_para = re.sub(r"\s+", " ", para_text)
            normalized_chunk = re.sub(r"\s+", " ", chunk_text).strip()

            found_at = normalized_para.find(normalized_chunk, page_search_start)
            if found_at == -1:
                # fallback: reset search start and try again
                found_at = normalized_para.find(normalized_chunk)
                if found_at == -1:
                    # last resort: omit exact offsets
                    chunk_start_char = None
                    chunk_end_char = None
                else:
                    chunk_start_char = found_at
                    chunk_end_char = found_at + len(normalized_chunk)
            else:
                chunk_start_char = found_at
                chunk_end_char = found_at + len(normalized_chunk)

            # Advance search start to keep subsequent matches stable
            if found_at != -1:
                page_search_start = found_at + max(1, len(normalized_chunk) // 3)

            chunk_id = stable_chunk_id(document_id, para_number if isinstance(para_number, int) else 0, idx, chunk_text)

            chunks.append({
                "document_id": document_id,
                "chunk_id": chunk_id,
                "text": chunk_text.strip(),
                "page_number": para_number if isinstance(para_number, int) else 0,
                "source_file": docx_path.name,
                "chunk_start_char": chunk_start_char,
                "chunk_end_char": chunk_end_char,
                "chunking_strategy": CHUNKING_STRATEGY,
                "created_at": created_at
            })

        global_char_cursor += len(para_text)

    return chunks


def write_jsonl(chunks, path: str):
    with open(path, "w", encoding="utf-8") as f:
        for obj in chunks:
            f.write(json.dumps(obj, ensure_ascii=False) + "\n")


def write_csv(chunks, path: str):
    # CSV-friendly column order
    fieldnames = [
        "document_id", "chunk_id", "page_number", "source_file",
        "chunk_start_char", "chunk_end_char",
        "chunking_strategy", "created_at", "text"
    ]
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for obj in chunks:
            w.writerow({k: obj.get(k) for k in fieldnames})


if __name__ == "__main__":
    import sys
    
    # Allow command line override of source file
    if len(sys.argv) > 1:
        source_file = sys.argv[1]
    else:
        source_file = SOURCE_FILE
    
    print(f"Processing DOCX: {source_file}")
    chunks = build_chunks(source_file)
    write_jsonl(chunks, OUT_JSONL)
    write_csv(chunks, OUT_CSV)
    print(f"Wrote {len(chunks)} chunks to {OUT_JSONL} and {OUT_CSV}")
